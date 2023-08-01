import os

import docx

def walkFile(file, doc):

    for root, dirs, files in os.walk(file):

        # root 表示当前正在访问的文件夹路径

        # dirs 表示该文件夹下的子目录名list

        # files 表示该文件夹下的文件list

        # 遍历文件

        for f in files:
            if f.endswith(".py") and f != "__init__.py":
                with open(os.path.join(root, f), "r", encoding="utf8") as f:
                    data = f.read()
                    doc.add_paragraph(data)


        # 遍历所有的文件夹

        for d in dirs:

            walkFile(d, doc)

doc = docx.Document()
# walkFile("./module/eq_mt_item", doc)
# walkFile("./module/ess_mt_item", doc)
# walkFile("./module/eq_mt_plan", doc)
# walkFile("./module/ess_mt_plan", doc)
# walkFile("./module/eq_mt_plan_timetable", doc)
# walkFile("./module/ess_mt_plan_timetable", doc)
# walkFile("./module/eq_mt_task", doc)
# walkFile("./module/ess_mt_task", doc)
# walkFile("./module/eq_mt_sub_task", doc)
# walkFile("./module/ess_mt_sub_task", doc)
######################
# walkFile("./module/event", doc)
# walkFile("./module/problem", doc)
# walkFile("./module/equipment", doc)
# walkFile("./model", doc)
###########################

# doc.save("DCOS事件模块.docx")


# doc = docx.Document()
# walkFile("./module/check_patrol_project", doc)
# walkFile("./module/check_region", doc)
# walkFile("./module/check_route", doc)
# walkFile("./module/check_task", doc)
# doc.save("DCOS运维巡检模块.docx")

################################
walkFile("./module/equipment", doc)
walkFile("./module/eq_information", doc)
walkFile("./module/dictionary", doc)
walkFile("./module/data_center", doc)
walkFile("./module/data_center_building", doc)
walkFile("./module/data_center_floor", doc)
walkFile("./module/equipment_basic_cockpit", doc)
walkFile("./module/equipment_sys", doc)

doc.save("DCOS设备信息模块.docx")




