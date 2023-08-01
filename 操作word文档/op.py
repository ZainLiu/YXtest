from io import BytesIO

import docx
import requests
from docx.shared import Pt

year = 2023
month = 7
day = 28
bc = "白班"
dutier = "陈焕鑫 王胜保"
path = "./temple_raw.docx"
new_path = "./templer_real.docx"
data = {
    "eng_room": [
        {
            "title": "1、各设备运行正常。",
            "value": 1
        },
        {
            "title": "2、动环、监控、BA系统运行正常。",
            "value": 0
        }
    ],
    "files": [
        {
            "title": "1、日常巡检表、水电表齐全，填写完整。",
            "value": 1
        },
        {
            "title": "2、交接班表齐全，填写完整。",
            "value": 1
        }
    ],
    "tools": [
        {
            "title": "1、工具仪表是否齐全。",
            "value": 1
        },
        {
            "title": "2、值班手机、对讲机是否完好。",
            "value": 1
        },
        {
            "title": "3、U盘、仓库应急卡、钥匙是否完好。",
            "value": 1
        }
    ]
}
docx_file = requests.get("https://opspre.imyunxia.com/file/2023_07_28/2023_07_28_14_23_10_81.docx")
doc = docx.Document(BytesIO(docx_file.content))
# for i, paragraph in enumerate(doc.paragraphs):
#     if i == 1:
#         print(paragraph.text)
#         temp_text = paragraph.text
#         temp_text = temp_text.replace("[1]", str(year))
#         temp_text = temp_text.replace("[2]", str(month))
#         temp_text = temp_text.replace("[3]", str(day))
#         temp_text = temp_text.replace("[4]", bc)
#         paragraph.text = temp_text
#
#         for run in paragraph.runs:
#             run.font.name = "宋体"
#             run.font.size = Pt(10)
#             run.font.bold = True

table = doc.tables[0]


# 值班人员
# dutier_cell = table.rows[0].cells[0]
# dutier_cell.text = dutier_cell.text.replace("[1]", dutier)
# for paragraph in dutier_cell.paragraphs:
#     for run in paragraph.runs:
#         run.font.name = "宋体"
#         run.font.size = Pt(7.5)
#         run.font.bold = True
# print(cell.text)
# cell.text = "2333333333333"

def confirm_items(table, row, col, data, title):
    tools_cell = table.rows[row].cells[col]
    tools_info_list = []
    for info in data:
        status_str = "（ √ ）" if info["value"] else "（ × ）"
        tools_info_list.append(info["title"] + status_str)

    tools_cell.text = f"{title}  " + '\t'.join(tools_info_list)
    for paragraph in tools_cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run.font.size = Pt(7.5)
            run.font.bold = True


# 机房运行情况交接
# eng_room_cell = table.rows[1].cells[0]
# eng_room_info_list = []
# for info in data["eng_room"]:
#     status_str = "（ √ ）" if info["value"] else "（ × ）"
#     eng_room_info_list.append(info["title"] + status_str)
#
# eng_room_cell.text = "机房运行情况交接:  " + '\t'.join(eng_room_info_list)
# for paragraph in eng_room_cell.paragraphs:
#     for run in paragraph.runs:
#         run.font.name = "宋体"
#         run.font.size = Pt(7.5)
#         run.font.bold = True

# 工具、器具情况交接
# tools_cell = table.rows[2].cells[0]
# tools_info_list = []
# for info in data["tools"]:
#     status_str = "（ √ ）" if info["value"] else "（ × ）"
#     tools_info_list.append(info["title"] + status_str)
#
# tools_cell.text = "工具、器具情况交接:  " + '\t'.join(tools_info_list)
# for paragraph in tools_cell.paragraphs:
#     for run in paragraph.runs:
#         run.font.name = "宋体"
#         run.font.size = Pt(7.5)
#         run.font.bold = True


# 文件情况交接
# confirm_items(table, 3, 0, data["files"], "文件情况交接")

# 总结
summary = """交接情况总结：
1.	当班交接班表齐全，填写完整
2.	当班工具、仪表数量交接完整
3.	当班事件已交接清楚
交班人确认：常亮 甄雨澎 林景俊1
接班人确认：马渊耀 潘文亮 陈奋
"""
summery_cell = table.rows[-1].cells[0]
summery_cell.text = summary

upload_url = "https://opspre.imyunxia.com/upload/manage/upload"
byte_io = BytesIO()
doc.save(byte_io)
resp = requests.post(upload_url, data={"project_id": 1}, files={"file": byte_io.getvalue()})
print(resp.json())
print()

doc.save(new_path)
# if __name__ == '__main__':
#     with open("template.docx","r") as f:
#         print(f.read())